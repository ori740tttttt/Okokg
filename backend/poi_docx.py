"""
Generazione e lettura di un documento WORD (.docx) per le attrazioni della mappa.

- build_poi_docx(pois): crea un .docx con una scheda-tabella per ogni attrazione.
  Nome e ubicazione sono precompilati; i campi info si compilano scrivendo nelle
  celle di destra. Facile da modificare in Word / Pages, anche su iPad.

- parse_poi_docx(raw): legge le tabelle e restituisce le attrazioni con i campi
  compilati: { match: 'id'|'new', id, fields:{...} } pronte per l'upsert.
"""
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Etichette riga -> chiave campo
ROW_DEFS = [
    ("ID (non modificare)", "id", True),       # readonly di fatto
    ("Nome", "name", False),
    ("Comune", "town", False),
    ("Provincia (PA/TP)", "province", False),
    ("Categoria (art/beach/nature)", "category", False),
    ("Coordinate (non modificare)", "_coords", True),
    ("Descrizione", "description", False),
    ("Prezzo (es. €6, ridotto €3)", "price", False),
    ("Orari (es. 09:00–19:30)", "hours", False),
    ("Durata (es. 2–3 ore)", "duration", False),
    ("Sconti (es. under 18 gratis)", "discount", False),
    ("Note utili", "notes", False),
    ("Link biglietti (https://…)", "ticket_url", False),
    ("Link foto (https://…)", "image_url", False),
]

# mappa etichetta-normalizzata -> chiave (per il parsing)
LABEL_TO_KEY = {label.split("(")[0].strip().lower(): key for label, key, _ in ROW_DEFS}

EDITABLE_KEYS = ["name", "town", "province", "category", "description",
                 "price", "hours", "duration", "discount", "notes",
                 "ticket_url", "image_url"]

TERRACOTTA = RGBColor(0xC2, 0x6B, 0x4E)
INK = RGBColor(0x2B, 0x26, 0x22)
INK_SOFT = RGBColor(0x6B, 0x62, 0x59)
SAND = "FBF7F1"
SAND_LINE = "E7DED2"


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_poi_docx(pois) -> bytes:
    doc = Document()

    # margini stretti per più spazio
    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    title = doc.add_paragraph()
    run = title.add_run("Appartamento Matteo — Schede attrazioni")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = TERRACOTTA

    intro = doc.add_paragraph()
    r = intro.add_run("Compila la colonna di destra di ogni scheda, poi salva e reimporta "
                      "il file Word nell'area gestione › Mappa interattiva. "
                      "Non modificare le righe ID e Coordinate.")
    r.font.size = Pt(9)
    r.font.color.rgb = INK_SOFT

    for poi in pois:
        # intestazione nome
        h = doc.add_paragraph()
        h.space_before = Pt(10)
        rn = h.add_run(poi.get("name", ""))
        rn.bold = True
        rn.font.size = Pt(12.5)
        rn.font.color.rgb = INK

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        # larghezze colonne
        for label, key, readonly in ROW_DEFS:
            row = table.add_row()
            lc, vc = row.cells[0], row.cells[1]
            lc.width = Inches(2.3)
            vc.width = Inches(4.7)

            # etichetta
            lp = lc.paragraphs[0]
            lr = lp.add_run(label)
            lr.font.size = Pt(9)
            lr.font.color.rgb = INK_SOFT if not readonly else RGBColor(0x9A, 0x90, 0x86)
            if readonly:
                lr.italic = True
            _shade_cell(lc, SAND)

            # valore
            if key == "_coords":
                lat, lng = poi.get("lat"), poi.get("lng")
                value = f"{lat}, {lng}" if lat is not None and lng is not None else ""
            else:
                value = poi.get(key)
                value = "" if value is None else str(value)
            vp = vc.paragraphs[0]
            vr = vp.add_run(value)
            vr.font.size = Pt(10)
            vr.font.color.rgb = INK
            if readonly:
                vr.font.color.rgb = RGBColor(0x9A, 0x90, 0x86)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def parse_poi_docx(raw: bytes) -> list:
    """Legge il .docx e restituisce una lista di record:
       { 'id': str|'', 'fields': {chiave: valore} } per ogni tabella/scheda."""
    doc = Document(BytesIO(raw))
    records = []
    for table in doc.tables:
        rec = {}
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.split("(")[0].strip().lower()
            value = row.cells[1].text.strip()
            key = LABEL_TO_KEY.get(label)
            if not key or key == "_coords":
                continue
            rec[key] = value
        # considera tabella valida solo se ha un id o un nome
        if rec.get("id") or rec.get("name"):
            records.append(rec)
    return records
